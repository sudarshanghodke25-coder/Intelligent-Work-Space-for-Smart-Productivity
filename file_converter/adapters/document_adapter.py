"""
file_converter/adapters/document_adapter.py
Handles DOCX/DOC ↔ text/HTML/Markdown conversions using python-docx.
PDF output for Word is delegated to the OS print-to-PDF pathway via
comtypes (Windows) or LibreOffice subprocess.
"""

from __future__ import annotations

import subprocess
import sys
from pathlib import Path

from file_converter.adapters.base_adapter import BaseAdapter
from file_converter.models.conversion_job import ConversionJob
from file_converter.exceptions.converter_errors import (
    CorruptFileError, ConversionFailureError, JobCancelledError,
)


class DocumentAdapter(BaseAdapter):
    """
    Converts Microsoft Word documents to/from other formats.
    Requires: python-docx
    Optional: comtypes (Windows) for .docx → .pdf via Word COM automation
    """

    adapter_name = "DocumentAdapter"
    supported_input_formats  = [".docx", ".doc", ".pptx", ".ppt", ".txt", ".md", ".html"]
    supported_output_formats = [".docx", ".txt", ".html", ".md", ".pdf"]

    def convert(self, job: ConversionJob) -> str:
        self._check_cancelled(job)
        src_ext = job.source_ext.lower()
        tgt_ext = job.target_ext.lower()
        job.report_progress(0.05, "Analysing document…")

        try:
            if src_ext in (".docx", ".doc") and tgt_ext == ".txt":
                return self._docx_to_text(job)
            elif src_ext in (".docx", ".doc") and tgt_ext == ".html":
                return self._docx_to_html(job)
            elif src_ext in (".docx", ".doc", ".pptx", ".ppt") and tgt_ext == ".pdf":
                return self._to_pdf_via_com(job)
            elif src_ext in (".docx", ".doc") and tgt_ext == ".md":
                return self._docx_to_markdown(job)
            elif src_ext == ".txt" and tgt_ext == ".docx":
                return self._text_to_docx(job)
            elif src_ext == ".md" and tgt_ext == ".docx":
                return self._markdown_to_docx(job)
            else:
                raise ConversionFailureError(
                    f"DocumentAdapter cannot handle {src_ext} → {tgt_ext}"
                )
        except JobCancelledError:
            raise
        except ConversionFailureError:
            raise
        except Exception as exc:
            raise ConversionFailureError(str(exc), exc) from exc

    # ── DOCX → TXT ─────────────────────────────────────────────────────────

    def _docx_to_text(self, job: ConversionJob) -> str:
        try:
            import docx
        except ImportError:
            raise ConversionFailureError("python-docx not installed.")
        job.report_progress(0.2, "Extracting text…")
        try:
            doc = docx.Document(job.source_path)
        except Exception as exc:
            raise CorruptFileError(str(exc), exc) from exc

        paragraphs = [p.text for p in doc.paragraphs]
        self._check_cancelled(job)
        job.report_progress(0.7, "Writing output…")
        out_path = self._resolve_output_path(job)
        out_path.write_text("\n".join(paragraphs), encoding="utf-8")
        job.report_progress(1.0, "Done!")
        return str(out_path)

    # ── DOCX → HTML ────────────────────────────────────────────────────────

    def _docx_to_html(self, job: ConversionJob) -> str:
        try:
            import docx
        except ImportError:
            raise ConversionFailureError("python-docx not installed.")
        job.report_progress(0.2, "Parsing document…")
        try:
            doc = docx.Document(job.source_path)
        except Exception as exc:
            raise CorruptFileError(str(exc), exc) from exc

        html_parts = ["<!DOCTYPE html><html><head><meta charset='utf-8'></head><body>"]
        for i, para in enumerate(doc.paragraphs):
            self._check_cancelled(job)
            style = para.style.name.lower()
            if "heading 1" in style:
                html_parts.append(f"<h1>{para.text}</h1>")
            elif "heading 2" in style:
                html_parts.append(f"<h2>{para.text}</h2>")
            elif "heading 3" in style:
                html_parts.append(f"<h3>{para.text}</h3>")
            else:
                html_parts.append(f"<p>{para.text}</p>")
            job.report_progress(0.2 + 0.6 * (i + 1) / max(len(doc.paragraphs), 1),
                                f"Converting paragraph {i + 1}…")

        html_parts.append("</body></html>")
        out_path = self._resolve_output_path(job)
        out_path.write_text("\n".join(html_parts), encoding="utf-8")
        job.report_progress(1.0, "Done!")
        return str(out_path)

    # ── DOCX → Markdown ────────────────────────────────────────────────────

    def _docx_to_markdown(self, job: ConversionJob) -> str:
        try:
            import docx
        except ImportError:
            raise ConversionFailureError("python-docx not installed.")
        job.report_progress(0.2, "Parsing document…")
        try:
            doc = docx.Document(job.source_path)
        except Exception as exc:
            raise CorruptFileError(str(exc), exc) from exc

        md_lines = []
        for para in doc.paragraphs:
            self._check_cancelled(job)
            style = para.style.name.lower()
            if "heading 1" in style:
                md_lines.append(f"# {para.text}")
            elif "heading 2" in style:
                md_lines.append(f"## {para.text}")
            elif "heading 3" in style:
                md_lines.append(f"### {para.text}")
            elif para.text.strip():
                md_lines.append(para.text)
            else:
                md_lines.append("")

        out_path = self._resolve_output_path(job)
        out_path.write_text("\n".join(md_lines), encoding="utf-8")
        job.report_progress(1.0, "Done!")
        return str(out_path)

    # ── Document → PDF (Word & PowerPoint) ─────────────────────────────────────────────────────────

    def _to_pdf_via_com(self, job: ConversionJob) -> str:
        """
        Attempt Office COM automation on Windows.
        Fallback: LibreOffice headless subprocess.
        """
        job.report_progress(0.1, "Converting to PDF…")
        self._check_cancelled(job)
        out_path = self._resolve_output_path(job)

        # Strategy 1: comtypes / Office COM (Windows only)
        if sys.platform == "win32":
            try:
                src_ext = job.source_ext.lower()
                if src_ext in (".pptx", ".ppt"):
                    return self._pptx_to_pdf_via_powerpoint_com(job, out_path)
                else:
                    return self._docx_to_pdf_via_word_com(job, out_path)
            except Exception:
                pass  # fall through to LibreOffice

        # Strategy 2: LibreOffice headless
        try:
            return self._docx_to_pdf_via_libreoffice(job, out_path)
        except Exception as exc:
            raise ConversionFailureError(
                f"Document→PDF requires Microsoft Office (Windows) or LibreOffice: {exc}", exc
            ) from exc

    def _pptx_to_pdf_via_powerpoint_com(self, job: ConversionJob, out_path: Path) -> str:
        """Use Microsoft PowerPoint COM automation (Windows only)."""
        try:
            import comtypes.client
        except ImportError:
            raise RuntimeError("comtypes not available.")

        ppt = None
        pres = None
        try:
            ppt = comtypes.client.CreateObject("Powerpoint.Application")
            ppt.DisplayAlerts = 1  # ppAlertsNone = 1
            # PowerPoint COM requires presentation to be opened. We don't make it visible if possible.
            pres = ppt.Presentations.Open(str(Path(job.source_path).resolve()), ReadOnly=True, WithWindow=False)
            job.report_progress(0.6, "Printing to PDF…")
            self._check_cancelled(job)
            pres.SaveAs(str(out_path.resolve()), 32)  # 32 = ppSaveAsPDF
            job.report_progress(1.0, "Done!")
            return str(out_path)
        finally:
            if pres is not None:
                try:
                    pres.Close()
                except Exception:
                    pass
            if ppt is not None:
                try:
                    if ppt.Presentations.Count == 0:
                        ppt.Quit()
                except Exception:
                    try:
                        ppt.Quit()
                    except Exception:
                        pass

    def _docx_to_pdf_via_word_com(self, job: ConversionJob, out_path: Path) -> str:
        """Use Microsoft Word COM automation (Windows only)."""
        try:
            import comtypes.client
        except ImportError:
            raise RuntimeError("comtypes not available.")

        word = None
        doc = None
        try:
            word = comtypes.client.CreateObject("Word.Application")
            word.Visible = False
            word.DisplayAlerts = 0  # wdAlertsNone
            doc = word.Documents.Open(str(Path(job.source_path).resolve()), ReadOnly=True)
            job.report_progress(0.6, "Printing to PDF…")
            self._check_cancelled(job)
            doc.ExportAsFixedFormat(OutputFileName=str(out_path.resolve()), ExportFormat=17)  # 17 = wdExportFormatPDF
            job.report_progress(1.0, "Done!")
            return str(out_path)
        finally:
            if doc is not None:
                try:
                    doc.Close(False)
                except Exception:
                    pass
            if word is not None:
                try:
                    word.Quit()
                except Exception:
                    pass

    def _docx_to_pdf_via_libreoffice(self, job: ConversionJob, out_path: Path) -> str:
        """Use LibreOffice headless conversion."""
        result = subprocess.run(
            ["soffice", "--headless", "--convert-to", "pdf",
             "--outdir", str(out_path.parent), job.source_path],
            capture_output=True, text=True, timeout=120,
        )
        if result.returncode != 0:
            raise RuntimeError(result.stderr)
        # LibreOffice saves as <stem>.pdf in the same directory
        expected = out_path.parent / (Path(job.source_path).stem + ".pdf")
        if expected.exists() and expected != out_path:
            expected.rename(out_path)
        job.report_progress(1.0, "Done!")
        return str(out_path)

    # ── TXT → DOCX ─────────────────────────────────────────────────────────

    def _text_to_docx(self, job: ConversionJob) -> str:
        try:
            import docx
        except ImportError:
            raise ConversionFailureError("python-docx not installed.")
        job.report_progress(0.2, "Reading text file…")
        try:
            text = Path(job.source_path).read_text(encoding="utf-8", errors="replace")
        except Exception as exc:
            raise CorruptFileError(str(exc), exc) from exc
        self._check_cancelled(job)
        job.report_progress(0.5, "Building Word document…")
        doc = docx.Document()
        for line in text.splitlines():
            doc.add_paragraph(line)
        out_path = self._resolve_output_path(job)
        doc.save(str(out_path))
        job.report_progress(1.0, "Done!")
        return str(out_path)

    # ── Markdown → DOCX ────────────────────────────────────────────────────

    def _markdown_to_docx(self, job: ConversionJob) -> str:
        try:
            import docx
        except ImportError:
            raise ConversionFailureError("python-docx not installed.")
        job.report_progress(0.2, "Parsing Markdown…")
        try:
            text = Path(job.source_path).read_text(encoding="utf-8", errors="replace")
        except Exception as exc:
            raise CorruptFileError(str(exc), exc) from exc
        self._check_cancelled(job)
        doc = docx.Document()
        for line in text.splitlines():
            if line.startswith("# "):
                doc.add_heading(line[2:], level=1)
            elif line.startswith("## "):
                doc.add_heading(line[3:], level=2)
            elif line.startswith("### "):
                doc.add_heading(line[4:], level=3)
            elif line.startswith("- ") or line.startswith("* "):
                doc.add_paragraph(line[2:], style="List Bullet")
            else:
                doc.add_paragraph(line)
        out_path = self._resolve_output_path(job)
        doc.save(str(out_path))
        job.report_progress(1.0, "Done!")
        return str(out_path)
