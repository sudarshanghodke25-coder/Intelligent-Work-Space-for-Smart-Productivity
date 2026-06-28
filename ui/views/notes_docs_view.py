import customtkinter as ctk
from tkinter import filedialog
import threading
import os
import shutil
import json
from theme import Colors, Fonts, Dims, blend_color
from ui.glass_card import GlassCard
from services.event_bus import bus
from services.knowledge_service import knowledge_service
from services.knowledge_pipeline import knowledge_pipeline
import webbrowser
import tkinter as tk

class NotesDocsView(ctk.CTkFrame):
    def __init__(self, parent, **kwargs):
        super().__init__(parent, fg_color="transparent", **kwargs)
        
        # Grid layout (2 columns now, left panel removed)
        self.grid_rowconfigure(0, weight=1)
        self.grid_columnconfigure(0, weight=1, minsize=400) # Center Panel (Sources list)
        self.grid_columnconfigure(1, weight=2, minsize=500) # Right Panel (AI Knowledge)
        
        # Center Panel (Search, List, Upload)
        self.center_panel = ctk.CTkFrame(self, fg_color="transparent")
        self.center_panel.grid(row=0, column=0, sticky="nsew", padx=(10, 10), pady=10)
        
        # Right Panel (AI Knowledge Panel)
        self.right_panel = GlassCard(self, title="")
        self.right_panel.grid(row=0, column=1, sticky="nsew", padx=(10, 4), pady=10)
        
        # State tracking
        self.current_source_id = None
        self.current_state = "BLANK" # BLANK, LOADING, FAILED, AI_VIEW
        
        self.filter_type = "All Sources"
        
        self._search_debounce_id = None
        self._refresh_pending = False
        self._selected_source_ids = set()
        
        # Track ONLY user-added content widgets inside doc_list.
        # NEVER call winfo_children() on a CTkScrollableFrame to clear it ---
        # that also returns its internal canvas/scrollbar, and destroying those
        # triggers Tcl_Panic (C-level abort that bypasses all Python exception handlers).
        self._doc_list_widgets = []
        
        self._build_center_panel()
        self._build_right_panel()
        
        # Subscribe to backend events
        bus.subscribe("ANALYSIS_STARTED", self._on_analysis_started)
        bus.subscribe("ANALYSIS_PROGRESS", self._on_analysis_progress)
        bus.subscribe("ANALYSIS_COMPLETED", self._on_analysis_completed)
        
        self._load_sources()

    def on_show(self):
        self._schedule_refresh()

    # ==========================================
    # CENTER PANEL
    # ==========================================
    def _build_center_panel(self):
        # Top Bar: Analyzers
        top_bar = ctk.CTkFrame(self.center_panel, fg_color="transparent")
        top_bar.pack(fill="x", pady=(0, 10))
        
        # 1. Document Analyzer
        doc_frame = ctk.CTkFrame(top_bar, fg_color=Colors.CARD_FLOATING, corner_radius=8, border_width=1, border_color=Colors.BORDER_SUBTLE)
        doc_frame.pack(fill="x", pady=(0, 5))
        
        ctk.CTkLabel(doc_frame, text="Document Analyzer", font=Fonts.SMALL_BOLD, text_color=Colors.TEXT_PRIMARY).pack(side="left", padx=10, pady=10)
        
        self.doc_path_lbl = ctk.CTkLabel(doc_frame, text="No file selected...", font=Fonts.ENTRY, text_color=Colors.TEXT_MUTED)
        self.doc_path_lbl.pack(side="left", fill="x", expand=True, padx=10)
        
        self.doc_path_val = None
        
        ctk.CTkButton(doc_frame, text="Upload File", font=Fonts.BUTTON, width=100, fg_color=blend_color(Colors.ACCENT_PRIMARY, 0.4), hover_color=blend_color(Colors.ACCENT_PRIMARY, 0.6), command=self._select_file).pack(side="left", padx=5)
        ctk.CTkButton(doc_frame, text="Analyze Document", font=Fonts.BUTTON, width=120, fg_color=Colors.ACCENT_PRIMARY, command=self._analyze_document).pack(side="right", padx=(5, 10))
        
        # 2. URL Analyzer
        url_frame = ctk.CTkFrame(top_bar, fg_color=Colors.CARD_FLOATING, corner_radius=8, border_width=1, border_color=Colors.BORDER_SUBTLE)
        url_frame.pack(fill="x", pady=(5, 10))
        
        ctk.CTkLabel(url_frame, text="URL Analyzer", font=Fonts.SMALL_BOLD, text_color=Colors.TEXT_PRIMARY).pack(side="left", padx=10, pady=10)
        
        self.url_input = ctk.CTkEntry(url_frame, placeholder_text="Enter URL...", font=Fonts.ENTRY, height=Dims.ENTRY_HEIGHT, fg_color=Colors.INPUT_BG, border_color=Colors.INPUT_BORDER)
        self.url_input.pack(side="left", fill="x", expand=True, padx=5)
        
        ctk.CTkButton(url_frame, text="Analyze URL", font=Fonts.BUTTON, width=100, fg_color=Colors.ACCENT_PRIMARY, command=self._analyze_url).pack(side="right", padx=(5, 10))
        
        # Search & Filter
        filter_row = ctk.CTkFrame(top_bar, fg_color="transparent")
        filter_row.pack(fill="x", pady=(10, 0))
        
        self.search_entry = ctk.CTkEntry(filter_row, placeholder_text="Search Sources...", font=Fonts.ENTRY, height=Dims.ENTRY_HEIGHT, fg_color=Colors.INPUT_BG, border_color=Colors.INPUT_BORDER)
        self.search_entry.pack(side="left", fill="x", expand=True, padx=(0, 10))
        self.search_entry.bind("<KeyRelease>", self._on_search_keyrelease)
        
        self.type_filter_btn = ctk.CTkOptionMenu(
            filter_row, values=["All Sources", "Documents", "Websites", "YouTube", "Notes"],
            font=Fonts.BUTTON, fg_color=Colors.CARD_FLOATING, button_color=Colors.CARD_HOVER,
            command=self._handle_type_filter
        )
        self.type_filter_btn.set("Source Type ▼")
        self.type_filter_btn.pack(side="left")
        
        # History Label & Bulk Actions
        hist_row = ctk.CTkFrame(self.center_panel, fg_color="transparent")
        hist_row.pack(fill="x", padx=10, pady=(15, 5))
        
        self.history_label = ctk.CTkLabel(hist_row, text="History (0)", font=Fonts.HEADING, text_color=Colors.TEXT_PRIMARY)
        self.history_label.pack(side="left")
        
        self.btn_clear_all = ctk.CTkButton(hist_row, text="Clear All", font=Fonts.SMALL, width=60, fg_color=Colors.ERROR, hover_color=blend_color(Colors.ERROR, 0.6), command=self._action_clear_all)
        self.btn_clear_all.pack(side="right", padx=5)
        
        self.btn_delete_selected = ctk.CTkButton(hist_row, text="Delete Selected", font=Fonts.SMALL, width=110, fg_color=Colors.ERROR, hover_color=blend_color(Colors.ERROR, 0.6), command=self._action_delete_selected)
        self.btn_delete_selected.pack(side="right", padx=5)
        
        self.btn_select_all = ctk.CTkButton(hist_row, text="Select All", font=Fonts.SMALL, width=70, fg_color=Colors.CARD_FLOATING, hover_color=Colors.CARD_HOVER, command=self._action_select_all)
        self.btn_select_all.pack(side="right", padx=5)
        
        # Source List Area
        self.doc_list = ctk.CTkScrollableFrame(self.center_panel, fg_color="transparent")
        self.doc_list.pack(fill="both", expand=True)

    def _handle_type_filter(self, filter_val: str):
        self.type_filter_btn.set(f"{filter_val} ▼")
        self.filter_type = filter_val
        self._schedule_refresh()

    def _on_search_keyrelease(self, event=None):
        if self._search_debounce_id:
            self.after_cancel(self._search_debounce_id)
        self._search_debounce_id = self.after(300, self._schedule_refresh)

    def _schedule_refresh(self):
        if self._refresh_pending:
            return
        self._refresh_pending = True
        self.after(50, self._do_refresh)

    def _do_refresh(self):
        self._refresh_pending = False
        self._load_sources()

    def _select_file(self):
        filepath = filedialog.askopenfilename()
        if not filepath: return
        self.doc_path_val = filepath
        self.doc_path_lbl.configure(text=os.path.basename(filepath))
        
    def _analyze_document(self):
        if not self.doc_path_val: return
        val = self.doc_path_val
        print(f"[UI ACTION] Analyze Document button fired. File: {val}")
        self.doc_path_val = None
        self.doc_path_lbl.configure(text="No file selected...")
        
        filename = os.path.basename(val)
        ext = os.path.splitext(val)[1].lower().replace('.', '')
        stype = "pdf" if ext == "pdf" else "document"
        if ext in ['txt', 'md', 'csv']: stype = "document" # generic doc
        
        upload_dir = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "..", "data", "uploads")
        os.makedirs(upload_dir, exist_ok=True)
        dest_path = os.path.join(upload_dir, filename)
        
        if os.path.exists(val):
            try:
                shutil.copy2(val, dest_path)
            except Exception:
                dest_path = val
        else:
            dest_path = val
            
        knowledge_pipeline.process_file_background(
            file_path=dest_path, title=filename, original_filename=filename,
            source_type=stype
        )
        
    def _analyze_url(self):
        val = self.url_input.get().strip()
        if not val: return
        
        source_type = "website"
        print(f"[UI ACTION] Analyze URL button fired. Type: {source_type}, URL: {val}")
        self.url_input.delete(0, "end")
        
        title = "Website Source"
        knowledge_pipeline.process_url_background(url=val, title=title, source_type=source_type)

    def _reextract_source(self):
        if not self.current_source_id: return
        source = knowledge_service.get_source_details(self.current_source_id)
        if not source: return
        
        # Trigger the pipeline again based on source type
        path = source.get("source_path")
        url = source.get("url")
        stype = source.get("source_type")
        title = source.get("title")
        
        title = source.get("title")
        
        if path:
            knowledge_pipeline.process_file_background(file_path=path, title=title, original_filename=source.get("original_filename"), source_type=stype)
        elif url:
            knowledge_pipeline.process_url_background(url=url, title=title, source_type=stype)

    def _reextract_specific_source(self, source):
        path = source.get("source_path")
        url = source.get("url")
        stype = source.get("source_type")
        title = source.get("title")
        
        if path:
            knowledge_pipeline.process_file_background(file_path=path, title=title, original_filename=source.get("original_filename"), source_type=stype)
        elif url:
            knowledge_pipeline.process_url_background(url=url, title=title, source_type=stype)

    def _get_icon_for_type(self, source_type: str) -> str:
        return {"document": "📄", "pdf": "📑", "youtube": "▶️", "website": "🌐", "note": "📝"}.get(source_type.lower(), "📄")
        
    def _get_color_for_type(self, source_type: str) -> str:
        return {"document": "#1E88E5", "pdf": "#E53935", "youtube": "#E53935", "website": "#43A047", "note": "#8E24AA"}.get(source_type.lower(), Colors.ACCENT_PRIMARY)

    def _load_sources(self, event=None):
        try:
            search_query = self.search_entry.get().strip().lower()
        except Exception:
            search_query = ""
        
        for widget in self._doc_list_widgets:
            try:
                if widget.winfo_exists():
                    widget.destroy()
            except Exception:
                pass
        self._doc_list_widgets.clear()
            
        filters = {}
        if search_query:
            filters["search"] = search_query
            
        # Map UI filter to DB types
        if self.filter_type == "Documents":
            filters["type"] = ["document", "pdf"]
        elif self.filter_type == "Websites":
            filters["type"] = ["website"]
        elif self.filter_type == "YouTube":
            filters["type"] = ["youtube"]
        elif self.filter_type == "Notes":
            filters["type"] = ["note"]
            
        sources = knowledge_service.get_sources(filters=filters if filters else None)
        
        if hasattr(self, 'history_label'):
            self.history_label.configure(text=f"History ({len(sources)})")
        
        if not sources:
            self._build_empty_state()
            return
            
        for s in sources:
            self._build_source_card(s)

    def _build_empty_state(self):
        f = ctk.CTkFrame(self.doc_list, fg_color="transparent")
        f.pack(fill="both", expand=True, pady=60)
        self._doc_list_widgets.append(f)  # Track for safe cleanup
        
        ctk.CTkLabel(f, text="No Knowledge Sources Yet", font=Fonts.HEADING, text_color=Colors.TEXT_PRIMARY).pack(pady=(0, 10))
        ctk.CTkLabel(f, text="Use the input bar above to analyze a Document, Website, YouTube video, or Note.", font=Fonts.BODY, text_color=Colors.TEXT_MUTED).pack(pady=(0, 20))

    def _build_source_card(self, source):
        is_selected = (source["id"] == self.current_source_id)
        bg_color = blend_color(Colors.CARD_HOVER, 0.2) if is_selected else Colors.CARD_FLOATING
        border_color = Colors.ACCENT_PRIMARY if is_selected else Colors.BORDER_SUBTLE
        
        card = ctk.CTkFrame(self.doc_list, fg_color=bg_color, corner_radius=12, border_width=1, border_color=border_color)
        card.pack(fill="x", pady=6, padx=2)
        setattr(card, "source_id", source["id"])
        self._doc_list_widgets.append(card)  # Track for safe cleanup (NOT via winfo_children())
        
        # Select Checkbox
        chk_var = ctk.StringVar(value="on" if source["id"] in self._selected_source_ids else "off")
        chk = ctk.CTkCheckBox(card, text="", variable=chk_var, onvalue="on", offvalue="off", width=24, command=lambda c=chk_var, sid=source["id"]: self._on_card_checkbox(sid, c.get()))
        chk.pack(side="left", padx=(10, 0))
        setattr(card, "chk_var", chk_var)
        
        # Click event binding on card itself
        card.bind("<Button-1>", lambda e, s=source: self._open_source(s))
            
        # Left Icon
        icon_frame = ctk.CTkFrame(card, fg_color=self._get_color_for_type(source.get("source_type", "document")), width=40, height=40, corner_radius=8)
        icon_frame.pack(side="left", padx=15, pady=15)
        icon_frame.pack_propagate(False)
        ctk.CTkLabel(icon_frame, text=self._get_icon_for_type(source.get("source_type", "document")), font=("Inter", 18)).pack(expand=True)
        
        icon_frame.bind("<Button-1>", lambda e, s=source: self._open_source(s))
        
        # Details
        details_frame = ctk.CTkFrame(card, fg_color="transparent")
        details_frame.pack(side="left", fill="both", expand=True, pady=15)
        details_frame.bind("<Button-1>", lambda e, s=source: self._open_source(s))

        
        title = source.get("title", "Untitled")
        title_lbl = ctk.CTkLabel(details_frame, text=title, font=Fonts.BODY_BOLD, text_color=Colors.TEXT_PRIMARY, anchor="w", justify="left", wraplength=280)
        title_lbl.pack(fill="x")
        title_lbl.bind("<Button-1>", lambda e, s=source: self._open_source(s))
        
        status = source.get("status", "COMPLETED")
        status_colors = {"COMPLETED": "🟢", "PROCESSING": "🟡", "FAILED": "🔴", "PENDING": "🟡"}
        status_icon = status_colors.get(status, "🟢")
        
        meta_text = f"{status_icon} {status.title()} • {source.get('updated_at', 'Unknown').split(' ')[0]}"
        if source.get('size'):
            meta_text += f" • {source.get('size')//1024} KB"
        meta_lbl = ctk.CTkLabel(details_frame, text=meta_text, font=Fonts.CAPTION, text_color=Colors.TEXT_MUTED, anchor="w", justify="left", wraplength=280)
        meta_lbl.pack(fill="x", pady=(2, 0))
        meta_lbl.bind("<Button-1>", lambda e, s=source: self._open_source(s))
        
        # Actions
        def open_card_menu():
            menu = tk.Menu(self, tearoff=0, bg=Colors.CARD_FLOATING, fg=Colors.TEXT_PRIMARY,
                           activebackground=Colors.ACCENT_PRIMARY, activeforeground="white",
                           font=(Fonts.BODY[0], 10))
            menu.add_command(label="Open", command=lambda: self._open_source(source))
            menu.add_command(label="Re-analyze", command=lambda: self._reextract_specific_source(source))
            menu.add_separator()
            menu.add_command(label="Delete", command=lambda: self._delete_source(source))
            
            x = self.winfo_pointerx()
            y = self.winfo_pointery()
            menu.tk_popup(x, y)
            
        ctk.CTkButton(card, text="⋮", font=Fonts.BODY, width=30, fg_color="transparent", hover_color=Colors.CARD_HOVER, command=open_card_menu).pack(side="right", padx=10)

    # ==========================================
    # RIGHT PANEL (AI KNOWLEDGE PANEL)
    # ==========================================
    def _build_right_panel(self):
        import traceback
        try:
            if not self.winfo_exists() or not self.right_panel.winfo_exists():
                print("[STEP] _build_right_panel: widget destroyed, skipping rebuild.")
                return
            for widget in self.right_panel.content.winfo_children():
                widget.destroy()
                
            if self.current_state == "BLANK":
                ctk.CTkLabel(self.right_panel.content, text="Select a source to view AI Knowledge Insights", font=Fonts.BODY, text_color=Colors.TEXT_MUTED).pack(expand=True)
                
            elif self.current_state == "LOADING":
                frame = ctk.CTkFrame(self.right_panel.content, fg_color="transparent")
                frame.pack(expand=True)
                ctk.CTkLabel(frame, text="⏳ Extracting Knowledge...", font=Fonts.HEADING, text_color=Colors.TEXT_PRIMARY).pack(pady=10)
                
                self.progress_lbl = ctk.CTkLabel(frame, text=getattr(self, 'current_progress', "Initializing..."), font=Fonts.BODY_BOLD, text_color=Colors.ACCENT_PRIMARY)
                self.progress_lbl.pack(pady=10)
                
                ctk.CTkLabel(frame, text="AUREX Intelligence Pipeline is actively analyzing the source.", font=Fonts.BODY, text_color=Colors.TEXT_MUTED).pack()
                
            elif self.current_state == "FAILED":
                print(f"[UI RENDER] Displaying FAILED state for source_id: {self.current_source_id}")
                source_data = knowledge_service.get_source_details(self.current_source_id)
                err_msg = getattr(self, 'last_error', None) or (source_data.get('raw_content', 'Unknown error.') if source_data else "Unknown error.")
                
                err_box = ctk.CTkTextbox(self.right_panel.content, font=Fonts.BODY, fg_color="transparent", text_color=Colors.ERROR, wrap="word")
                err_box.pack(fill="both", expand=True, padx=20, pady=20)
                err_box.insert("1.0", err_msg)
                err_box.configure(state="disabled")
                
            elif self.current_state == "AI_VIEW":
                print(f"[UI RENDER] Reading analysis from database for source_id: {self.current_source_id}")
                source_data = knowledge_service.get_source_details(self.current_source_id)
                if not source_data:
                    return
                
                # Add Title explicitly inside content
                ctk.CTkLabel(self.right_panel.content, text=source_data.get('title', ''), font=Fonts.HEADING, text_color=Colors.TEXT_PRIMARY, anchor="w").pack(fill="x", pady=(0, 5))
                
                # Header Meta
                header_meta = ctk.CTkFrame(self.right_panel.content, fg_color="transparent")
                header_meta.pack(fill="x", pady=(0, 15))
                
                kb = f"{source_data.get('size', 0)//1024} KB • " if source_data.get('size') else ""
                
                proc_time = source_data.get('processing_time')
                duration_str = f" in {proc_time:.1f}s" if proc_time else ""
                
                meta_str = f"{source_data.get('source_type', 'document').upper()} • {kb}Analyzed{duration_str} on {source_data.get('updated_at', '').split(' ')[0]}"
                ctk.CTkLabel(header_meta, text=meta_str, font=Fonts.SMALL, text_color=Colors.TEXT_MUTED).pack(side="left")
                
                # Action Buttons
                def open_more_menu():
                    menu = tk.Menu(self, tearoff=0, bg=Colors.CARD_FLOATING, fg=Colors.TEXT_PRIMARY,
                                   activebackground=Colors.ACCENT_PRIMARY, activeforeground="white",
                                   font=(Fonts.BODY[0], 10))
                    
                    menu.add_command(label="Export Summary (.txt)", command=lambda: self._action_export_txt(source_data))
                    menu.add_command(label="Export Summary (.pdf)", command=lambda: self._action_export_pdf(source_data))
                    menu.add_command(label="Export Summary (.docx)", command=lambda: self._action_export_docx(source_data))
                    menu.add_command(label="Export Summary (.md)", command=lambda: self._action_export_md(source_data))
                    menu.add_command(label="View Raw Extracted Text", command=lambda: self._action_view_raw(source_data))
                    menu.add_separator()
                    menu.add_command(label="Regenerate Analysis", command=lambda: self._reextract_source())
                    menu.add_command(label="Rename Source", command=lambda: self._action_rename_source(source_data))
                    menu.add_separator()
                    menu.add_command(label="Delete Source", command=lambda: self._delete_source(source_data))
                    
                    x = self.winfo_pointerx()
                    y = self.winfo_pointery()
                    menu.tk_popup(x, y)

                btn_more = ctk.CTkButton(header_meta, text="⋮ More Actions", font=Fonts.BODY_BOLD, fg_color="transparent", hover_color=Colors.CARD_HOVER, width=100, command=open_more_menu)
                btn_more.pack(side="right", padx=2)
                
                btn_copy = ctk.CTkButton(header_meta, text="📋 Copy Analysis", font=Fonts.BODY_BOLD, fg_color="transparent", hover_color=Colors.CARD_HOVER, width=100, command=lambda: self._action_copy_analysis(source_data))
                btn_copy.pack(side="right", padx=2)
                
                btn_open = ctk.CTkButton(header_meta, text="🔗 Open Source", font=Fonts.BODY_BOLD, fg_color="transparent", hover_color=Colors.CARD_HOVER, width=100, command=lambda: self._action_open_source(source_data))
                btn_open.pack(side="right", padx=2)
                
                # Content Area (Tabs)
                self.tabs = ctk.CTkTabview(self.right_panel.content, fg_color="transparent")
                self.tabs.pack(fill="both", expand=True, pady=(10, 0))
                
                tab_summary = self.tabs.add("Summary")
                tab_transcript = self.tabs.add("Transcript")
                tab_chapters = self.tabs.add("Chapters")
                tab_concepts = self.tabs.add("Concepts")
                tab_quotes = self.tabs.add("Quotes")
                
                def add_section_to_tab(tab, title, content, height, fallback=""):
                    ctk.CTkLabel(tab, text=title, font=Fonts.BODY_BOLD, text_color=Colors.ACCENT_PRIMARY).pack(anchor="w", pady=(5, 5))
                    box = ctk.CTkTextbox(tab, font=Fonts.BODY, fg_color="transparent", text_color=Colors.TEXT_PRIMARY, wrap="word", height=height)
                    box.pack(fill="x", expand=False, pady=(0, 10))
                    
                    if content and content.strip():
                        box.insert("1.0", content)
                    else:
                        box.insert("1.0", fallback)
                    box.configure(state="disabled")
                
                # Summary Tab
                sum_scroll = ctk.CTkScrollableFrame(tab_summary, fg_color="transparent")
                sum_scroll.pack(fill="both", expand=True)
                add_section_to_tab(sum_scroll, "✨ Summary", source_data.get('summary'), 100, "No summary stored.")
                add_section_to_tab(sum_scroll, "🔑 Key Points", source_data.get('key_points'), 150, "No key points.")
                add_section_to_tab(sum_scroll, "🎯 Action Items", source_data.get('action_items'), 150, "No action items.")
                
                # Transcript Tab
                tr_scroll = ctk.CTkScrollableFrame(tab_transcript, fg_color="transparent")
                tr_scroll.pack(fill="both", expand=True)
                
                full_transcript = source_data.get('transcript', '') or ""
                if len(full_transcript) > 15000:
                    display_transcript = full_transcript[:15000] + "\n\n... [TRANSCRIPT TRUNCATED FOR UI PERFORMANCE. USE 'MORE ACTIONS -> EXPORT' TO VIEW FULL TEXT.]"
                else:
                    display_transcript = full_transcript
                    
                add_section_to_tab(tr_scroll, "📜 Transcript", display_transcript, 500, "No transcript stored.")
                
                # Chapters Tab
                ch_scroll = ctk.CTkScrollableFrame(tab_chapters, fg_color="transparent")
                ch_scroll.pack(fill="both", expand=True)
                add_section_to_tab(ch_scroll, "⏱️ Chapters", source_data.get('chapters'), 400, "No chapters available.")
                
                # Concepts Tab
                co_scroll = ctk.CTkScrollableFrame(tab_concepts, fg_color="transparent")
                co_scroll.pack(fill="both", expand=True)
                add_section_to_tab(co_scroll, "🧠 Important Concepts", source_data.get('important_concepts'), 150, "No concepts.")
                add_section_to_tab(co_scroll, "🏷️ Topics", source_data.get('topics'), 100, "No topics.")
                add_section_to_tab(co_scroll, "🏢 Entities Mentioned", f"People:\n{source_data.get('people_mentioned', '')}\n\nCompanies:\n{source_data.get('companies_mentioned', '')}", 150, "No entities.")
                
                # Quotes Tab
                qu_scroll = ctk.CTkScrollableFrame(tab_quotes, fg_color="transparent")
                qu_scroll.pack(fill="both", expand=True)
                add_section_to_tab(qu_scroll, "💬 Quotes", source_data.get('quotes'), 400, "No quotes stored.")

        except Exception:
            print("[CRITICAL] Exception in _build_right_panel:")
            traceback.print_exc()

    # ==========================================
    # LOGIC & EVENTS
    # ==========================================
    def _open_source(self, source):
        source_id = source["id"]
        
        # Load directly from database
        source_data = knowledge_service.get_source_details(source_id)
        if not source_data:
            return
            
        print("========== SOURCE DATA ==========")
        print("ID:", source_data.get('id'))
        print("TITLE:", source_data.get('title'))
        print("STATUS:", source_data.get('status'))
        print("SUMMARY LENGTH:", len(source_data.get('summary') or ""))
        print("KEY POINTS LENGTH:", len(source_data.get('key_points') or ""))
        print("SIMPLE EXPLANATION LENGTH:", len(source_data.get('simple_explanation') or ""))
        print("TRANSCRIPT LENGTH:", len(source_data.get('transcript') or ""))
        print("IMPORTANT CONCEPTS LENGTH:", len(source_data.get('important_concepts') or ""))
        print("================================")
        
        self.current_source_id = source_id
        status = source_data.get('status', 'COMPLETED')
        
        if status in ['PENDING', 'PROCESSING']:
            self.current_state = "LOADING"
            self.current_progress = "Extracting Knowledge..."
        elif status == 'FAILED':
            self.current_state = "FAILED"
            self.last_error = source_data.get('raw_content', "Analysis failed.")
        else:
            self.current_state = "AI_VIEW"
            
        self._build_right_panel()
        self._update_card_selection()

    def _update_card_selection(self):
        for widget in self._doc_list_widgets:
            if hasattr(widget, "source_id") and widget.winfo_exists():
                is_selected = (widget.source_id == self.current_source_id)
                bg_color = blend_color(Colors.CARD_HOVER, 0.2) if is_selected else Colors.CARD_FLOATING
                border_color = Colors.ACCENT_PRIMARY if is_selected else Colors.BORDER_SUBTLE
                widget.configure(border_color=border_color, fg_color=bg_color)

    def _delete_source(self, source):
        knowledge_service.delete_source(source["id"])
        if self.current_source_id == source["id"]:
            self.current_state = "BLANK"
            self.current_source_id = None
            self._build_right_panel()
        self._schedule_refresh()

    def _on_card_checkbox(self, source_id, val):
        if val == "on":
            self._selected_source_ids.add(source_id)
        else:
            self._selected_source_ids.discard(source_id)

    def _action_select_all(self):
        # Toggle based on current state
        all_selected = all(getattr(w, "chk_var", ctk.StringVar(value="off")).get() == "on" for w in self._doc_list_widgets if hasattr(w, "chk_var"))
        new_val = "off" if all_selected else "on"
        
        for widget in self._doc_list_widgets:
            if hasattr(widget, "chk_var"):
                widget.chk_var.set(new_val)
                self._on_card_checkbox(widget.source_id, new_val)

    def _action_delete_selected(self):
        if not self._selected_source_ids: return
        knowledge_service.delete_sources(list(self._selected_source_ids))
        if self.current_source_id in self._selected_source_ids:
            self.current_state = "BLANK"
            self.current_source_id = None
            self._build_right_panel()
        self._selected_source_ids.clear()
        self._schedule_refresh()

    def _action_clear_all(self):
        # Confirmation dialog
        dialog = ctk.CTkToplevel(self)
        dialog.title("Confirm Clear All")
        dialog.geometry("400x150")
        dialog.transient(self)
        dialog.grab_set()
        
        ctk.CTkLabel(dialog, text="Are you sure you want to delete ALL knowledge sources?\nThis cannot be undone.", font=Fonts.BODY, text_color=Colors.TEXT_PRIMARY).pack(pady=20)
        
        btn_frame = ctk.CTkFrame(dialog, fg_color="transparent")
        btn_frame.pack(pady=10)
        
        def confirm():
            knowledge_service.delete_all_sources()
            self.current_state = "BLANK"
            self.current_source_id = None
            self._selected_source_ids.clear()
            self._build_right_panel()
            self._schedule_refresh()
            dialog.destroy()
            
        ctk.CTkButton(btn_frame, text="Cancel", width=80, fg_color=Colors.CARD_FLOATING, hover_color=Colors.CARD_HOVER, command=dialog.destroy).pack(side="left", padx=10)
        ctk.CTkButton(btn_frame, text="Delete All", width=80, fg_color=Colors.ERROR, hover_color=blend_color(Colors.ERROR, 0.6), command=confirm).pack(side="left", padx=10)

    def _on_analysis_started(self, payload):
        # bus.publish already dispatches on main thread via app.after(0) — no need for another self.after(0)
        print(f"[STEP] _on_analysis_started received. source_id={payload.get('source_id')}")
        self._handle_analysis_started(payload)

    def _handle_analysis_started(self, payload):
        import traceback
        try:
            if not self.winfo_exists():
                return
            self._schedule_refresh()
            if payload.get("source_id") == self.current_source_id or not self.current_source_id:
                self.current_source_id = payload.get("source_id")
                self.current_state = "LOADING"
                self.current_progress = "Validating..."
                self.last_error = None
                self._build_right_panel()
        except Exception:
            print("[CRITICAL] Exception in _handle_analysis_started:")
            traceback.print_exc()

    def _on_analysis_progress(self, payload):
        # bus.publish already dispatches on main thread via app.after(0) — no need for another self.after(0)
        self._handle_analysis_progress(payload)
        
    def _handle_analysis_progress(self, payload):
        if payload.get("source_id") == self.current_source_id:
            self.current_progress = payload.get("status", "Processing...")
            if hasattr(self, 'progress_lbl') and self.progress_lbl.winfo_exists():
                self.progress_lbl.configure(text=self.current_progress)

    def _on_analysis_completed(self, payload):
        # NOTE: bus.publish already marshals via app.after(0, ...) so calling
        # self.after(0) here creates a DOUBLE-DISPATCH — the callback is already
        # on the main thread. We can call the handler directly.
        print(f"[STEP] _on_analysis_completed received. source_id={payload.get('source_id')}, success={payload.get('success')}")
        self._handle_analysis_completed(payload)

    def _handle_analysis_completed(self, payload):
        import traceback
        try:
            print("[STEP] _handle_analysis_completed entered")
            if not self.winfo_exists():
                print("[STEP] _handle_analysis_completed: widget no longer exists, skipping.")
                return
            print("[STEP] Before _load_sources")
            self._schedule_refresh()
            print("[STEP] After _load_sources")
            if payload.get("source_id") == self.current_source_id:
                if payload.get("success"):
                    self.current_state = "AI_VIEW"
                else:
                    self.current_state = "FAILED"
                    self.last_error = payload.get("error", "Unknown error.")
                print("[STEP] Before _build_right_panel (state={self.current_state})")
                self._build_right_panel()
                print("[STEP] After _build_right_panel")
        except Exception:
            print("[CRITICAL] Exception in _handle_analysis_completed:")
            traceback.print_exc()

    # ==========================================
    # ACTION METHODS
    # ==========================================
    def _action_open_source(self, source_data):
        path = source_data.get('source_path')
        url = source_data.get('url')
        try:
            if url:
                webbrowser.open(url)
            elif path and os.path.exists(path):
                os.startfile(path)
        except Exception as e:
            print(f"Error opening source: {e}")

    def _action_copy_analysis(self, source_data):
        summary = source_data.get('simple_explanation', '')
        points = source_data.get('important_concepts', '')
        text = f"Summary:\n{summary}\n\nKey Points & Concepts:\n{points}"
        self.clipboard_clear()
        self.clipboard_append(text)

    def _action_export_txt(self, source_data):
        title = source_data.get('title', 'Export').replace(" ", "_")
        path = filedialog.asksaveasfilename(defaultextension=".txt", initialfile=f"{title}_Analysis.txt", filetypes=[("Text Files", "*.txt")])
        if path:
            summary = source_data.get('simple_explanation', '')
            points = source_data.get('important_concepts', '')
            with open(path, "w", encoding="utf-8") as f:
                f.write(f"Title: {source_data.get('title', '')}\n")
                f.write("="*40 + "\n\n")
                f.write("SUMMARY:\n")
                f.write(summary + "\n\n")
                f.write("KEY POINTS & CONCEPTS:\n")
                f.write(points + "\n")

    def _action_export_pdf(self, source_data):
        title = source_data.get('title', 'Export').replace(" ", "_")
        path = filedialog.asksaveasfilename(defaultextension=".pdf", initialfile=f"{title}_Analysis.pdf", filetypes=[("PDF Files", "*.pdf")])
        if path:
            try:
                from reportlab.pdfgen import canvas
                from reportlab.lib.pagesizes import letter
                import textwrap
                
                c = canvas.Canvas(path, pagesize=letter)
                width, height = letter
                
                c.setFont("Helvetica-Bold", 16)
                c.drawString(50, height - 50, source_data.get('title', ''))
                
                c.setFont("Helvetica", 12)
                y = height - 80
                
                def write_text_block(header, text_content):
                    nonlocal y
                    c.setFont("Helvetica-Bold", 12)
                    c.drawString(50, y, header)
                    y -= 20
                    c.setFont("Helvetica", 10)
                    for line in text_content.split('\n'):
                        wrapped = textwrap.wrap(line, width=90)
                        if not wrapped:
                            y -= 15
                        for wline in wrapped:
                            c.drawString(50, y, wline)
                            y -= 15
                            if y < 50:
                                c.showPage()
                                c.setFont("Helvetica", 10)
                                y = height - 50
                    y -= 20

                write_text_block("Summary", source_data.get('simple_explanation', ''))
                write_text_block("Key Points & Concepts", source_data.get('important_concepts', ''))
                
                c.save()
            except Exception as e:
                print(f"Failed to export PDF: {e}")

    def _action_export_docx(self, source_data):
        title = source_data.get('title', 'Export').replace(" ", "_")
        path = filedialog.asksaveasfilename(defaultextension=".docx", initialfile=f"{title}_Analysis.docx", filetypes=[("Word Document", "*.docx")])
        if path:
            try:
                import docx
                doc = docx.Document()
                doc.add_heading(source_data.get('title', 'Media Analysis'), 0)
                
                def add_section(header, content):
                    if content and content.strip():
                        doc.add_heading(header, level=1)
                        doc.add_paragraph(content)

                add_section('Summary', source_data.get('summary', ''))
                add_section('Key Points', source_data.get('key_points', ''))
                add_section('Action Items', source_data.get('action_items', ''))
                add_section('Important Concepts', source_data.get('important_concepts', ''))
                add_section('Chapters', source_data.get('chapters', ''))
                add_section('Quotes', source_data.get('quotes', ''))
                
                doc.save(path)
            except Exception as e:
                print(f"Failed to export DOCX: {e}")

    def _action_export_md(self, source_data):
        title = source_data.get('title', 'Export').replace(" ", "_")
        path = filedialog.asksaveasfilename(defaultextension=".md", initialfile=f"{title}_Analysis.md", filetypes=[("Markdown", "*.md")])
        if path:
            try:
                with open(path, "w", encoding="utf-8") as f:
                    f.write(f"# {source_data.get('title', 'Media Analysis')}\n\n")
                    
                    def add_section(header, content):
                        if content and content.strip():
                            f.write(f"## {header}\n\n{content}\n\n")

                    add_section('Summary', source_data.get('summary', ''))
                    add_section('Key Points', source_data.get('key_points', ''))
                    add_section('Action Items', source_data.get('action_items', ''))
                    add_section('Important Concepts', source_data.get('important_concepts', ''))
                    add_section('Chapters', source_data.get('chapters', ''))
                    add_section('Quotes', source_data.get('quotes', ''))
            except Exception as e:
                print(f"Failed to export Markdown: {e}")

    def _action_view_raw(self, source_data):
        win = ctk.CTkToplevel(self)
        win.title("Raw Extracted Content")
        win.geometry("800x600")
        box = ctk.CTkTextbox(win, font=Fonts.BODY, wrap="word")
        box.pack(fill="both", expand=True, padx=20, pady=20)
        box.insert("1.0", source_data.get("raw_content", "No raw content stored."))
        box.configure(state="disabled")

    def _action_rename_source(self, source_data):
        dialog = ctk.CTkInputDialog(text="Enter new title:", title="Rename Source")
        new_title = dialog.get_input()
        if new_title and new_title.strip():
            knowledge_service.update_source_metadata(source_data['id'], {"title": new_title.strip()})
            self.right_panel.title_label.configure(text=new_title.strip()) if hasattr(self.right_panel, 'title_label') else None
            self._schedule_refresh()
            self._build_right_panel()


